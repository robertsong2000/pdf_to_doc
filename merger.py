"""
DOCX文件合并核心模块 - 使用 docxcompose 实现
"""
from docx import Document
from docxcompose.composer import Composer
import os
from typing import List


class DocxMerger:
    """DOCX文件合并器 - 基于 docxcompose"""

    def __init__(self, output_file: str, first_file: str = None):
        """
        初始化合并器

        Args:
            output_file: 输出文件路径
            first_file: 第一个文档路径（可选），如果提供则以此为基础文档
        """
        self.output_file = output_file

        # 如果提供了第一个文档，直接用它作为基础，避免空白页问题
        if first_file and os.path.exists(first_file):
            master = Document(first_file)
            self.composer = Composer(master)
            self.merged_count = 1
        else:
            # 创建空文档作为基础
            master = Document()
            self.composer = Composer(master)
            self.merged_count = 0

    def merge_document(self, source_file: str, add_page_break_before: bool = True):
        """
        合并一个源文档到目标文档

        Args:
            source_file: 源文档路径
            add_page_break_before: 是否在文档前添加分页符
        """
        if not os.path.exists(source_file):
            raise FileNotFoundError(f"文件不存在: {source_file}")

        # 如果需要，在合并之前添加分页符
        # 这样分页符会出现在两个文档之间
        if add_page_break_before:
            # 添加一个包含分页符的段落
            self.composer.doc.add_page_break()

        # 读取要追加的文档
        doc_to_append = Document(source_file)

        # 使用 docxcompose 的 append 方法进行智能合并
        # docxcompose 会自动处理：
        # - 样式冲突（通过重命名冲突样式）
        # - 图片和嵌入资源
        # - 列表编号（可能不完全完美）
        # - 页眉页脚
        self.composer.append(
            doc_to_append,
            remove_property_fields=True  # 移除文档属性字段避免冲突
        )

        self.merged_count += 1

    def merge_documents(self, source_files: List[str], add_page_break: bool = True):
        """
        批量合并多个文档

        Args:
            source_files: 源文档路径列表
            add_page_break: 是否在文档之间添加分页符
        """
        total = len(source_files)

        for i, source_file in enumerate(source_files):
            # 跳过第一个文档（因为它已经作为基础文档了）
            if self.merged_count > 0 and i == 0:
                print(f"[{i + 1}/{total}] 跳过（已作为基础文档）: {os.path.basename(source_file)}")
                continue

            # 对于所有要合并的文档，都需要在前面加分页符
            # 这样每个文档都会在新的一页开始
            print(f"[{i + 1}/{total}] 正在合并: {os.path.basename(source_file)}")

            try:
                self.merge_document(
                    source_file,
                    add_page_break_before=add_page_break  # 每个文档前加分页符
                )
            except Exception as e:
                print(f"  警告: 合并 {os.path.basename(source_file)} 时出现问题: {e}")
                raise

    def save(self):
        """保存合并后的文档"""
        self.composer.save(self.output_file)
        print(f"\n成功! 已合并 {self.merged_count} 个文档到: {self.output_file}")
